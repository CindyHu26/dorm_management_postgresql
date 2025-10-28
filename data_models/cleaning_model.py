# data_models/cleaning_model.py

import pandas as pd
import json
from datetime import date
from dateutil.relativedelta import relativedelta
import database
from io import BytesIO
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx 未安裝，無法使用產生 Word 公告功能。")

# --- 常數 ---
SIMPLE_CLEANING = "宿舍簡易清掃"
MAJOR_CLEANING = "宿舍大掃除"
# --- 新增：固定月份 ---
SIMPLE_CLEANING_MONTHS_FIXED = [3, 9]
MAJOR_CLEANING_MONTHS_FIXED = [6, 12]

CLEANING_TYPES_INFO = {
    SIMPLE_CLEANING: {"frequency": 3, "months": SIMPLE_CLEANING_MONTHS_FIXED}, # frequency 仍保留供顯示
    MAJOR_CLEANING: {"frequency": 6, "months": MAJOR_CLEANING_MONTHS_FIXED}
}

# --- 輔助函數：計算下一個固定排程日期 ---
def _calculate_next_fixed_schedule(current_date: date, target_months: list) -> date:
    """
    從 current_date 開始，計算下一個落在 target_months 中的月份，並回傳該月的第一天。
    """
    current_year = current_date.year
    current_month = current_date.month

    next_date = None
    # 檢查當年剩餘月份
    for month in sorted(target_months):
        if month >= current_month:
            # 找到當年下一個目標月份
            next_date = date(current_year, month, 1)
            break

    # 如果當年找不到，則找明年的第一個目標月份
    if next_date is None:
        first_target_month = min(target_months)
        next_date = date(current_year + 1, first_target_month, 1)

    return next_date

def _execute_query_to_dataframe(conn, query, params=None):
    """輔助函式，用來手動執行查詢並回傳 DataFrame。"""
    with conn.cursor() as cursor:
        cursor.execute(query, params)
        records = cursor.fetchall()
        if not records:
            columns = [desc[0] for desc in cursor.description] if cursor.description else []
            return pd.DataFrame([], columns=columns)
        columns = [desc[0] for desc in cursor.description]
        return pd.DataFrame(records, columns=columns)

def _get_my_company_dorm_ids(cursor):
    """輔助函式，取得由「我司」管理的宿舍 ID 列表。"""
    cursor.execute('SELECT id FROM "Dormitories" WHERE primary_manager = %s', ('我司',))
    return [row['id'] for row in cursor.fetchall()]

def initialize_cleaning_schedule(dorm_id: int):
    """
    【固定月份版】如果指定宿舍 ID 尚無清掃排程紀錄，則為其初始化（簡易和大掃除）。
    根據今天計算 *下一個* 固定的排程日期。
    如果初始化成功則回傳 True，否則 False。
    """
    conn = database.get_db_connection()
    if not conn:
        print("ERROR: 無法連接資料庫進行初始化。")
        return False

    initialized = False
    try:
        with conn.cursor() as cursor:
            today = date.today()
            for record_type, info in CLEANING_TYPES_INFO.items():
                # 檢查紀錄是否已存在
                cursor.execute(
                    'SELECT id FROM "ComplianceRecords" WHERE dorm_id = %s AND record_type = %s',
                    (dorm_id, record_type)
                )
                if cursor.fetchone():
                    continue # 如果已存在則跳過

                # 計算下一個固定排程日期
                next_schedule_date = _calculate_next_fixed_schedule(today, info["months"])

                details_dict = {
                    "last_completion_date": None, # 初始為 None
                    "next_schedule_date": next_schedule_date.strftime('%Y-%m-%d'),
                    "frequency_months": info["frequency"] # 保留供顯示
                }
                details_json = json.dumps(details_dict, ensure_ascii=False)

                cursor.execute(
                    'INSERT INTO "ComplianceRecords" (dorm_id, record_type, details) VALUES (%s, %s, %s)',
                    (dorm_id, record_type, details_json)
                )
                initialized = True
        conn.commit()
        return initialized
    except Exception as e:
        print(f"初始化宿舍 {dorm_id} 清掃排程時發生錯誤: {e}")
        if conn: conn.rollback()
        return False
    finally:
        if conn: conn.close()

def batch_initialize_schedules():
    """
    檢查所有由「我司」管理的宿舍，並為缺少排程的宿舍進行初始化。
    回傳新初始化的宿舍數量。
    """
    conn = database.get_db_connection()
    if not conn:
        print("ERROR: 無法連接資料庫進行批次初始化。")
        return 0

    newly_initialized_count = 0
    try:
        with conn.cursor() as cursor:
            my_dorm_ids = _get_my_company_dorm_ids(cursor)
            if not my_dorm_ids:
                return 0

            # 找出至少有一筆清掃紀錄的宿舍
            cursor.execute(
                'SELECT DISTINCT dorm_id FROM "ComplianceRecords" WHERE record_type = ANY(%s) AND dorm_id = ANY(%s)',
                (list(CLEANING_TYPES_INFO.keys()), my_dorm_ids)
            )
            dorms_with_schedule = {row['dorm_id'] for row in cursor.fetchall()}

            # 篩選出需要初始化的宿舍
            dorms_to_initialize = [dorm_id for dorm_id in my_dorm_ids if dorm_id not in dorms_with_schedule]

        # 逐一呼叫初始化函數（內部會處理連線）
        for dorm_id in dorms_to_initialize:
            if initialize_cleaning_schedule(dorm_id): # 使用更新後的日期計算邏輯
                 newly_initialized_count += 1 # 計算實際初始化的宿舍數

        return newly_initialized_count

    except Exception as e:
        print(f"批次初始化時發生錯誤: {e}")
        return newly_initialized_count # 回傳目前已處理的數量
    finally:
        if conn: conn.close()


def force_initialize_all_schedules(start_calculation_date: date):
    """
    【固定月份版】刪除所有「我司」宿舍現有的清掃排程，並根據提供的
    start_calculation_date 計算 *下一個* 固定的排程日期來建立新排程。
    回傳處理的宿舍數量和結果訊息。
    """
    conn = database.get_db_connection()
    if not conn:
        return 0, "資料庫連線失敗。"

    processed_count = 0
    errors = []
    try:
        with conn.cursor() as cursor:
            my_dorm_ids = _get_my_company_dorm_ids(cursor)
            if not my_dorm_ids:
                return 0, "找不到任何由「我司」管理的宿舍。"

            # 步驟 1: 刪除所有現有清掃排程
            cursor.execute(
                'DELETE FROM "ComplianceRecords" WHERE record_type = ANY(%s) AND dorm_id = ANY(%s)',
                (list(CLEANING_TYPES_INFO.keys()), my_dorm_ids)
            )
            deleted_count = cursor.rowcount
            print(f"INFO: 已刪除 {deleted_count} 筆現有的清掃排程紀錄。")

            # 步驟 2: 為每個宿舍重新建立排程
            for dorm_id in my_dorm_ids:
                try:
                    for record_type, info in CLEANING_TYPES_INFO.items():
                        # 根據起算日計算下一個固定排程
                        next_schedule_date = _calculate_next_fixed_schedule(start_calculation_date, info["months"])

                        details_dict = {
                            "last_completion_date": None, # 強制初始化
                            "next_schedule_date": next_schedule_date.strftime('%Y-%m-%d'),
                            "frequency_months": info["frequency"] # 保留供顯示
                        }
                        details_json = json.dumps(details_dict, ensure_ascii=False)

                        cursor.execute(
                            'INSERT INTO "ComplianceRecords" (dorm_id, record_type, details) VALUES (%s, %s, %s)',
                            (dorm_id, record_type, details_json)
                        )
                    processed_count += 1
                except Exception as dorm_error:
                    errors.append(f"宿舍 ID {dorm_id}: {dorm_error}")
                    # 記錄錯誤，但繼續處理下一個宿舍

            if errors:
                # 如果有任何錯誤，回滾整個交易
                 raise Exception("部分宿舍處理失敗，已復原所有操作。")

        # 只有在全部成功時才提交
        conn.commit()
        return processed_count, f"已成功為 {processed_count} 間宿舍根據 {start_calculation_date.strftime('%Y-%m-%d')} 重新設定清掃排程。"

    except Exception as e:
        if conn: conn.rollback()
        error_summary = "; ".join(errors) if errors else str(e)
        return 0, f"批次重設排程失敗: {error_summary}"
    finally:
        if conn: conn.close()

def clear_all_cleaning_schedules():
    """
    刪除所有由「我司」管理的宿舍的所有清掃排程紀錄。
    回傳刪除的紀錄數量和結果訊息。
    """
    conn = database.get_db_connection()
    if not conn:
        return 0, "資料庫連線失敗。"

    deleted_count = 0
    try:
        with conn.cursor() as cursor:
            my_dorm_ids = _get_my_company_dorm_ids(cursor)
            if not my_dorm_ids:
                return 0, "找不到任何由「我司」管理的宿舍。"

            cursor.execute(
                'DELETE FROM "ComplianceRecords" WHERE record_type = ANY(%s) AND dorm_id = ANY(%s)',
                (list(CLEANING_TYPES_INFO.keys()), my_dorm_ids)
            )
            deleted_count = cursor.rowcount
        conn.commit()
        return deleted_count, f"已成功清除 {deleted_count} 筆清掃排程紀錄。"
    except Exception as e:
        if conn: conn.rollback()
        return 0, f"清除排程時發生錯誤: {e}"
    finally:
        if conn: conn.close()


def get_cleaning_schedule(dorm_ids=None):
    """
    查詢指定宿舍 ID (或所有 '我司' 宿舍) 的目前清掃排程狀態。
    包含宿舍地址、縣市、區域、負責人等詳細資訊。
    """
    conn = database.get_db_connection()
    if not conn: return pd.DataFrame()

    try:
        target_dorm_ids = dorm_ids
        if target_dorm_ids is None:
            with conn.cursor() as cursor:
                 target_dorm_ids = _get_my_company_dorm_ids(cursor)
            if not target_dorm_ids:
                 return pd.DataFrame() # 找不到 '我司' 宿舍

        query = """
            SELECT
                cr.id,
                d.original_address AS "宿舍地址",
                d.city AS "縣市",
                d.district AS "區域",
                d.person_in_charge AS "負責人",
                cr.record_type AS "清掃類型",
                cr.details ->> 'last_completion_date' AS "上次完成日期",
                cr.details ->> 'next_schedule_date' AS "下次預計日期",
                (cr.details ->> 'frequency_months')::int AS "頻率(月)" -- 仍然查詢頻率供顯示
            FROM "ComplianceRecords" cr
            JOIN "Dormitories" d ON cr.dorm_id = d.id
            WHERE cr.record_type = ANY(%s) -- 篩選清掃類型
              AND cr.dorm_id = ANY(%s) -- 篩選宿舍 ID
            ORDER BY d.original_address, cr.record_type;
        """
        params = (list(CLEANING_TYPES_INFO.keys()), target_dorm_ids)
        df = _execute_query_to_dataframe(conn, query, params)

        # 將日期字串轉為 date 物件
        if not df.empty:
            df['上次完成日期'] = pd.to_datetime(df['上次完成日期'], errors='coerce').dt.date
            df['下次預計日期'] = pd.to_datetime(df['下次預計日期'], errors='coerce').dt.date

        return df

    except Exception as e:
        print(f"查詢清掃排程時發生錯誤: {e}")
        return pd.DataFrame()
    finally:
        if conn: conn.close()

def mark_cleaning_complete(compliance_record_ids: list, completion_date=date.today()):
    """
    【固定月份版】將一或多筆清掃紀錄標記為在指定日期完成，
    並計算 *下一個* 固定的排程日期。在單一交易中完成。
    """
    if not compliance_record_ids:
        return False, "未選擇任何要標記完成的紀錄。"

    conn = database.get_db_connection()
    if not conn: return False, "資料庫連線失敗。"

    updated_count = 0
    errors = []

    try:
        with conn.cursor() as cursor:
            for record_id in compliance_record_ids:
                try:
                    # 鎖定紀錄並取得類型和詳情
                    cursor.execute(
                        'SELECT record_type, details FROM "ComplianceRecords" WHERE id = %s FOR UPDATE',
                        (record_id,)
                    )
                    record = cursor.fetchone()
                    if not record or not record.get('details') or not record.get('record_type'):
                        errors.append(f"ID {record_id}: 找不到紀錄或缺少必要資料。")
                        continue

                    details_dict = record['details']
                    record_type = record['record_type']

                    # 根據類型找到目標月份
                    if record_type not in CLEANING_TYPES_INFO:
                         errors.append(f"ID {record_id}: 未知的清掃類型 ({record_type})。")
                         continue
                    target_months = CLEANING_TYPES_INFO[record_type]["months"]

                    # 計算下一個固定排程日期 (基準是完成日期的隔天)
                    next_schedule_date = _calculate_next_fixed_schedule(completion_date + relativedelta(days=1), target_months)

                    # 更新 details
                    details_dict['last_completion_date'] = completion_date.strftime('%Y-%m-%d')
                    details_dict['next_schedule_date'] = next_schedule_date.strftime('%Y-%m-%d')
                    # frequency_months 保留不變
                    details_json = json.dumps(details_dict, ensure_ascii=False)

                    # 執行更新
                    cursor.execute(
                        'UPDATE "ComplianceRecords" SET details = %s WHERE id = %s',
                        (details_json, record_id)
                    )
                    updated_count += 1
                except Exception as item_error:
                    errors.append(f"ID {record_id}: 更新失敗 - {item_error}")

            if errors:
                 # 若有任何錯誤，拋出異常以觸發 rollback
                 raise Exception("部分紀錄更新失敗，已復原所有操作。")

        # 所有紀錄都成功更新才提交
        conn.commit()
        return True, f"成功標記 {updated_count} 筆清掃紀錄完成，下次日期已更新。"

    except Exception as e:
        if conn: conn.rollback() # 確保錯誤時回滾
        error_summary = "; ".join(errors) if errors else str(e)
        return False, f"標記完成時發生錯誤: {error_summary}"
    finally:
        if conn: conn.close()


def batch_delete_cleaning_schedules(compliance_record_ids: list):
    """
    根據提供的 ID 列表，刪除特定的清掃排程紀錄。
    """
    if not compliance_record_ids:
        return 0, "未選擇任何要刪除的紀錄。"

    conn = database.get_db_connection()
    if not conn:
        return 0, "資料庫連線失敗。"

    deleted_count = 0
    try:
        with conn.cursor() as cursor:
            query = """
                DELETE FROM "ComplianceRecords"
                WHERE id = ANY(%s) AND record_type = ANY(%s)
            """
            cursor.execute(query, (compliance_record_ids, list(CLEANING_TYPES_INFO.keys())))
            deleted_count = cursor.rowcount
        conn.commit()
        return deleted_count, f"已成功刪除 {deleted_count} 筆清掃排程紀錄。"
    except Exception as e:
        if conn: conn.rollback()
        return 0, f"刪除排程時發生錯誤: {e}"
    finally:
        if conn: conn.close()


def generate_cleaning_notice(dorm_address: str, cleaning_date: date):
    """
    產生清掃公告 Word 文件 (BytesIO buffer)。
    需要安裝 python-docx 函式庫。
    """
    if not DOCX_AVAILABLE:
        raise ImportError("請先安裝 python-docx 函式庫以產生 Word 文件。")
    
    logo_path = "logo.png"

    date_str = f"{cleaning_date.month}/{cleaning_date.day}"
    time_str = "08:00"

    notice_text = {
        "zh": [ f"預計於{date_str} {time_str} 大掃除，", "請工人務必空出當天時間，並配合翻譯及服務人員指導，將宿舍環境打掃乾淨，如有自己私人物品要併入公共垃圾，要平均分攤清潔費用，當天個人物品請記得務必收好，若遺失概不負責。", ],
        "vi": [ f"Dự kiến sẽ tổng vệ sinh vào lúc {time_str} ngày {date_str}.", "Yêu cầu công nhân phải dành thời gian vào ngày hôm đó, phối hợp với phiên dịch và nhân viên hướng dẫn để làm sạch khu ký túc xá. Nếu có đồ dùng cá nhân muốn bỏ chung vào rác công cộng, cần chia đều chi phí vệ sinh. Vào ngày hôm đó, vui lòng tự bảo quản đồ cá nhân, nếu bị mất sẽ không chịu trách nhiệm.", ],
        "tl": [ f"Magkakaroon ng general cleaning sa {date_str} ng {time_str}.", "Ang lahat ng manggagawa ay inaasahang maglaan ng oras sa araw na iyon at makipag-cooperate sa mga tagapagsalin at mga tagapag-ugnay ng serbisyo upang malinis nang maayos ang buong dormitoryo. Kung may mga personal na gamit na isasama sa mga pampublikong basura, ang gastos sa paglilinis ay kailangang pantay-pantay na paghahati-hatian. Paalala: Siguraduhing maayos na itago ang mga personal na gamit sa araw ng paglilinis. Hindi kami mananagot sa anumang bagay na mawawala.", ],
        "th": [ f"กำหนดทำความสะอาดครั้งใหญ่ในวันที่ {date_str} เวลา {time_str}", "ขอให้คนงานทุกคนว่างในวันดังกล่าว และให้ความร่วมมือกับล่ามและเจ้าหน้าที่ ในการทำความสะอาดบริเวณหอพักให้เรียบร้อย หากมีของใช้ส่วนตัวที่จะทิ้งรวมกับขยะส่วนกลาง จะต้องแบ่งค่าใช้จ่ายในการทำความสะอาดอย่างเท่าเทียมกัน ในวันนั้น กรุณาเก็บของใช้ส่วนตัวให้เรียบร้อย หากสูญหาย ทางเราจะไม่รับผิดชอบใด ๆ ทั้งสิ้น", ],
        "id": [ f"Akan diadakan kerja bakti besar pada tanggal {date_str} pukul {time_str}.", "Para pekerja diwajibkan untuk meluangkan waktu pada hari tersebut dan bekerja sama dengan penerjemah serta petugas layanan untuk membersihkan lingkungan asrama dengan baik. Jika ada barang pribadi yang ingin dibuang bersama sampah umum, biaya kebersihan harus dibagi rata. Harap simpan barang pribadi dengan baik pada hari itu, jika hilang tidak akan menjadi tanggung jawab kami.", ]
    }

    try:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial Unicode MS' # 嘗試使用更廣泛支援 Unicode 的字型
        font.size = Pt(12)

        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)  # 段落前的間距設為 0 點
        paragraph_format.space_after = Pt(2)   # 段落後的間距設為 6 點 (您可以調整此數值，例如 Pt(3) 或 Pt(0))
        # 加入公告文字 (同前)
        for lang_code, lines in notice_text.items():
            for line in lines:
                 p = document.add_paragraph()
                 p.add_run(line)
            # document.add_paragraph() # 語言之間加空行

        # --- 在頁尾加入 Logo ---
        footer = document.sections[0].footer # 取得第一個節 (通常是唯一一個) 的頁尾
        # 如果頁尾有預設內容，可以先清除 (可選)
        # for para in footer.paragraphs:
        #     para.clear() # 清除段落內容
        # 如果需要完全移除段落:
        # if footer.paragraphs:
        #     # Iterate backwards when removing elements from a list
        #     for i in range(len(footer.paragraphs) - 1, -1, -1):
        #         p = footer.paragraphs[i]
        #         p._element.getparent().remove(p._element)

        # 在頁尾新增一個段落
        footer_para = footer.add_paragraph()
        # 設定段落置中
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # 在段落中加入圖片
        run = footer_para.add_run()
        try:
            # 嘗試加入圖片，可以設定寬度 (例如 1.5 英吋)
            run.add_picture(logo_path, width=Inches(3))
        except FileNotFoundError:
            print(f"錯誤：找不到 Logo 圖片檔案於 '{logo_path}'。頁尾將不包含圖片。")
            run.add_text("[Logo Image Not Found]") # 加入提示文字
        except Exception as img_err:
            print(f"錯誤：加入 Logo 圖片時發生問題: {img_err}")
            run.add_text("[Error Adding Logo]")

        # 儲存到 BytesIO buffer
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        print(f"產生 Word 公告時發生錯誤: {e}")
        return None