import streamlit as st
import pandas as pd
import numpy as np
from datetime import date
from data_models import maintenance_model, dormitory_model, vendor_model, equipment_model
import os
import io
import re
import base64  # 新增：用於圖片編碼
from urllib.parse import quote, unquote # 新增 unquote

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
except ImportError:
    pass

# 用於高效取得所有維修紀錄
@st.cache_data
def get_all_logs_for_selection():
    return maintenance_model.get_logs_for_view(filters=None)

# -----------------------------------------------------------------------------
# 輔助函式：圖片轉 Base64 (用於 HTML 報表)
# -----------------------------------------------------------------------------
def image_to_base64(image_path):
    """將本地圖片轉為 HTML 可用的 Base64 字串"""
    if not os.path.exists(image_path):
        return None
    try:
        with open(image_path, "rb") as img_file:
            encoded = base64.b64encode(img_file.read()).decode()
            # 判斷副檔名
            ext = os.path.splitext(image_path)[1].lower().replace('.', '')
            if ext == 'jpg': ext = 'jpeg'
            return f"data:image/{ext};base64,{encoded}"
    except Exception:
        return None

# -----------------------------------------------------------------------------
# 輔助函式：設定 Word 中文字型 (避免亂碼或字型跑掉)
# -----------------------------------------------------------------------------
def set_cell_font(cell, text, font_name='Microsoft JhengHei', font_size=10, bold=False):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text) if text else "")
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # 設定中文字型
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# -----------------------------------------------------------------------------
# 輔助函式：從說明中提取房號
# -----------------------------------------------------------------------------
def extract_room_number(description):
    """從 '【房號: 201】 說明...' 格式中提取 201"""
    if not isinstance(description, str):
        return ""
    match = re.search(r"【房號:\s*(.+?)】", description)
    if match:
        return match.group(1)
    return ""

# -----------------------------------------------------------------------------
# 共用元件：完整編輯表單 (封裝後可供「進度追蹤」與「編輯紀錄」共用)
# -----------------------------------------------------------------------------
def _render_full_edit_form(selected_log_id, dorm_options, vendor_options, item_type_options, status_options, key_suffix=""):
    """
    渲染單筆維修紀錄的完整編輯表單 (包含檔案管理、設備連動等)
    key_suffix: 用於區分不同頁面呼叫時的元件 key，避免衝突
    """
    details = maintenance_model.get_single_log_details(selected_log_id)
    if not details:
        st.error("找不到該筆資料，可能已被刪除。")
        return

    st.markdown(f"#### 正在編輯案件 ID: {selected_log_id}")

    # --- 1. 顯示既有檔案 ---
    st.markdown("##### 📂 檔案管理")
    existing_files = details.get('photo_paths') or []
    
# 顯示圖片預覽 (【修改】改用 columns 網格排列，解決照片黏在一起的問題)
    valid_images = [f for f in existing_files if os.path.exists(f) and f.lower().endswith(('.png', '.jpg', '.jpeg'))]
    
    if valid_images:
        st.write("###### 🖼️ 現場照片預覽：")
        # 設定每行顯示 4 張照片
        cols_per_row = 4
        cols = st.columns(cols_per_row)
        
        for idx, img_path in enumerate(valid_images):
            # 透過餘數運算決定放在第幾個欄位 (會自動換行)
            with cols[idx % cols_per_row]:
                st.image(
                    img_path, 
                    width=150, # 保持縮圖大小
                    caption=os.path.basename(img_path)
                )
    
    # 顯示 PDF 下載
    if pdf_files := [f for f in existing_files if os.path.exists(f) and f.lower().endswith('.pdf')]:
        st.write("📄 PDF 文件：")
        for pdf_path in pdf_files:
            with open(pdf_path, "rb") as pdf_file:
                st.download_button(
                    label=f"⬇️ 下載 {os.path.basename(pdf_path)}", 
                    data=pdf_file, 
                    file_name=os.path.basename(pdf_path), 
                    key=f"dl_{pdf_path}_{key_suffix}"
                )

    # --- 2. 編輯表單 ---
    with st.form(f"edit_log_form_{selected_log_id}_{key_suffix}"):
        
        # --- (A) 檔案刪除與上傳 ---
        files_to_delete = st.multiselect(
            "🗑️ 勾選要刪除的舊檔案：", 
            options=existing_files, 
            format_func=lambda f: os.path.basename(f),
            key=f"del_files_{selected_log_id}_{key_suffix}"
        )
        new_files = st.file_uploader(
            "📤 上傳新檔案 (圖片/PDF)", 
            type=['jpg', 'jpeg', 'png', 'pdf'], 
            accept_multiple_files=True,
            key=f"new_files_{selected_log_id}_{key_suffix}"
        )
        
        st.markdown("---")
        st.subheader("📋 案件資訊")
        ec1, ec2, ec3, ec4 = st.columns(4)
        
        # 宿舍處理
        current_dorm_id = details.get('dorm_id')
        dorm_keys = list(dorm_options.keys())
        current_dorm_index = dorm_keys.index(current_dorm_id) if current_dorm_id in dorm_keys else 0
        e_dorm_id = ec1.selectbox("宿舍地址", options=dorm_keys, format_func=lambda x: dorm_options.get(x, "未知"), index=current_dorm_index)
        
        # 設備處理 (根據當前選擇的宿舍，重新抓取設備選單)
        # 注意：Streamlit form 內無法做動態連動(選宿舍立刻變設備)，只能基於載入時的資料
        # 若要完全動態，需要把這兩格搬出 form，但為了版面整潔，我們先以此方式實作
        equipment_in_dorm_edit = equipment_model.get_equipment_for_view({"dorm_id": current_dorm_id}) if current_dorm_id else pd.DataFrame()
        equip_options_edit = {row['id']: f"{row['設備名稱']} ({row.get('位置', 'N/A')})" for _, row in equipment_in_dorm_edit.iterrows()}
        current_equip_id = details.get('equipment_id')
        
        equip_keys_list = [None] + list(equip_options_edit.keys())
        try:
            equip_index = equip_keys_list.index(current_equip_id)
        except ValueError:
            equip_index = 0

        e_equipment_id = ec2.selectbox("關聯設備", options=equip_keys_list, format_func=lambda x: "無" if x is None else equip_options_edit.get(x), index=equip_index)
        e_notification_date = ec3.date_input("收到通知日期", value=details.get('notification_date'))
        e_reported_by = ec4.text_input("公司內部提報人", value=details.get('reported_by'))
        
        st.subheader("🔧 維修詳情")
        edc1, edc2 = st.columns(2)

        current_item_type = details.get('item_type', '')
        if current_item_type in item_type_options:
            default_index = item_type_options.index(current_item_type)
            custom_val = ""
        else:
            default_index = item_type_options.index("其他(手動輸入)") if "其他(手動輸入)" in item_type_options else 0
            custom_val = current_item_type

        e_selected_item_type = edc1.selectbox("項目類型", options=item_type_options, index=default_index)
        e_custom_item_type = edc1.text_input("自訂項目類型 (若選其他)", value=custom_val)
        e_description = edc2.text_area("修理細項說明", value=details.get('description'), height=100)
        
        st.subheader("🏗️ 廠商與進度")
        ec6, ec7, ec8 = st.columns(3)
        status_idx = status_options.index(details.get('status')) if details.get('status') in status_options else 0
        e_status = ec6.selectbox("案件狀態", options=status_options, index=status_idx)
        
        vendor_keys = [None] + list(vendor_options.keys())
        vendor_idx = vendor_keys.index(details.get('vendor_id')) if details.get('vendor_id') in vendor_keys else 0
        e_vendor_id = ec7.selectbox("維修廠商", options=vendor_keys, format_func=lambda x: "未指定" if x is None else vendor_options.get(x), index=vendor_idx)
        e_contacted_vendor_date = ec7.date_input("聯絡廠商日期", value=details.get('contacted_vendor_date'))
        
        with ec8:
            e_completion_date = st.date_input("廠商回報完成日期", value=details.get('completion_date'))
        
        e_key_info = st.text_input("鑰匙/備註", value=details.get('key_info', ''))

        st.subheader("💰 費用與款項")
        ec9, ec10, ec11, ec12 = st.columns(4)
        e_cost = ec9.number_input("維修費用", min_value=0, step=100, value=details.get('cost') or 0)
        
        payer_opts = ["", "我司", "工人", "雇主", "房東"]
        payer_idx = payer_opts.index(details.get('payer')) if details.get('payer') in payer_opts else 0
        e_payer = ec10.selectbox("付款人", payer_opts, index=payer_idx)
        e_invoice_date = ec11.date_input("請款日期", value=details.get('invoice_date'))
        e_invoice_info = ec12.text_input("發票資訊", value=details.get('invoice_info', ''))

        e_notes = st.text_area("其他備註", value=details.get('notes'))

        # --- 儲存按鈕 ---
        if st.form_submit_button("💾 儲存完整變更"):
            final_type = e_custom_item_type if e_selected_item_type == "其他(手動輸入)" else e_selected_item_type
            
            # 處理檔案
            final_file_paths = [p for p in existing_files if p not in files_to_delete]
            if new_files:
                file_info_dict = {
                    "date": e_notification_date.strftime('%Y%m%d'), 
                    "address": dorm_options.get(e_dorm_id, 'UnknownAddr'), 
                    "reporter": e_reported_by, 
                    "type": final_type
                }
                for file in new_files:
                    path = maintenance_model.save_uploaded_photo(file, file_info_dict)
                    final_file_paths.append(path)

            update_data = {
                'dorm_id': e_dorm_id, 'equipment_id': e_equipment_id, 'status': e_status, 
                'vendor_id': e_vendor_id, 'notification_date': e_notification_date,
                'reported_by': e_reported_by, 'item_type': final_type, 'description': e_description,
                'contacted_vendor_date': e_contacted_vendor_date, 'completion_date': e_completion_date,
                'key_info': e_key_info, 'cost': e_cost, 'payer': e_payer, 'invoice_date': e_invoice_date,
                'invoice_info': e_invoice_info, 'notes': e_notes, 'photo_paths': final_file_paths 
            }
            
            success, message = maintenance_model.update_log(selected_log_id, update_data, paths_to_delete=files_to_delete)
            if success:
                st.success(f"儲存成功！ {message}")
                st.cache_data.clear()
                st.rerun()
            else:
                st.error(message)

# -----------------------------------------------------------------------------
# 子功能渲染函式
# -----------------------------------------------------------------------------

def render_add_new_record(dorm_options, vendor_options, item_type_options, status_options):
    """渲染：新增維修紀錄"""
    st.subheader("➕ 新增維修紀錄")

    if "maint_success_msg" in st.session_state:
        st.success(st.session_state.maint_success_msg)
        del st.session_state["maint_success_msg"]

    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        dorm_keys = list(dorm_options.keys())
        new_dorm_id = st.selectbox("宿舍 (連動設備)*", options=dorm_keys, format_func=lambda x: dorm_options.get(x, "未選擇"), key="add_m_dorm")

    with c2:
        if new_dorm_id:
            equipment_in_dorm = equipment_model.get_equipment_for_view({"dorm_id": new_dorm_id})
            if not equipment_in_dorm.empty:
                equip_options_new = {row['id']: f"{row['設備名稱']} ({row.get('位置', 'N/A')})" for _, row in equipment_in_dorm.iterrows()}
                new_equipment_id = st.selectbox("關聯設備 (選填)", options=[None] + list(equip_options_new.keys()), format_func=lambda x: "無" if x is None else equip_options_new.get(x), key="add_m_equip")
            else:
                st.selectbox("關聯設備", options=["該宿舍無設備資料"], disabled=True, key="add_m_equip_fake")
                new_equipment_id = None
        else:
            new_equipment_id = None

    with c3: new_report_date = st.date_input("收到通知日期*", value=date.today(), key="add_m_date")
    with c4: new_status = st.selectbox("案件狀態*", options=status_options, key="add_m_status")
    with c5:
        new_category_sel = st.selectbox("維修類別", options=item_type_options, key="add_m_cat")
        custom_category = st.text_input("請輸入自訂類型*", placeholder="例如: 網路費", key="add_m_cat_custom") if new_category_sel == "其他(手動輸入)" else None

    # 【修改點 1】 改為 6 個欄位 (c6~c11)，加入聯繫廠商日期
    c6, c7, c8, c9, c10, c11 = st.columns(6)
    with c6: new_cost = st.number_input("維修費用", min_value=0, step=100, key="add_m_cost")
    with c7: new_vendor = st.selectbox("廠商", options=[None]+list(vendor_options.keys()), format_func=lambda x: "未指定" if x is None else vendor_options.get(x), key="add_m_vendor")
    
    # 【新增】聯繫廠商日期
    with c8: new_contact_date = st.date_input("聯繫廠商日期", value=None, key="add_m_contact_date")
    
    with c9: new_payer = st.selectbox("付款人", ["", "我司", "工人", "雇主"], key="add_m_payer")
    with c10: new_finish_date = st.date_input("完成日期", value=None, key="add_m_finish")
    with c11: 
        st.write(""); st.write("")
        new_is_paid_check = st.checkbox("已付款?", value=False, key="add_m_paid_check")

    new_description = st.text_area("修理細項說明* (可換行)", height=150, placeholder="請詳細描述故障情形...", key="add_m_desc")
    
    c_sub1, c_sub2, c_sub3, c_sub4 = st.columns(4)
    new_reporter = c_sub1.text_input("提報人", placeholder="內部人員", key="add_m_reporter")
    new_key_info = c_sub2.text_input("鑰匙資訊", placeholder="如:警衛室", key="add_m_key_info")
    new_invoice_info = c_sub3.text_input("發票資訊", placeholder="抬頭/統編", key="add_m_invoice")
    new_notes = c_sub4.text_input("其他備註", placeholder="其他事項", key="add_m_notes")

    uploaded_files = st.file_uploader("📷 上傳照片/報價單 (可多選)", type=['jpg', 'jpeg', 'png', 'pdf'], accept_multiple_files=True, key="add_m_uploader")

    if st.button("💾 儲存維修案件", type="primary", use_container_width=True):
        final_category = custom_category if new_category_sel == "其他(手動輸入)" else new_category_sel
        if not new_dorm_id or not new_description:
            st.error("「宿舍」和「修理細項說明」為必填欄位！")
        elif new_category_sel == "其他(手動輸入)" and not custom_category:
            st.error("請填寫自訂類型名稱！")
        else:
            file_paths = []
            if uploaded_files:
                file_info_dict = {"date": new_report_date.strftime('%Y%m%d'), "address": dorm_options.get(new_dorm_id, 'UnknownAddr'), "reporter": new_reporter, "type": final_category}
                for file in uploaded_files:
                    path = maintenance_model.save_uploaded_photo(file, file_info_dict)
                    file_paths.append(path)
            
            final_status = new_status
            if new_finish_date and new_status in ["待處理", "待尋廠商", "進行中"]: final_status = "待付款"
            
            details = {
                'dorm_id': new_dorm_id, 'equipment_id': new_equipment_id, 'vendor_id': new_vendor, 'status': final_status,
                'notification_date': new_report_date, 'reported_by': new_reporter, 'item_type': final_category, 'description': new_description,
                # 【修改點 2】 將 None 改為 new_contact_date
                'contacted_vendor_date': new_contact_date, 
                'completion_date': new_finish_date, 'key_info': new_key_info,    
                'cost': new_cost, 'payer': new_payer, 'invoice_date': None, 'invoice_info': new_invoice_info, 
                'notes': new_notes, 'photo_paths': file_paths 
            }
            success, message = maintenance_model.add_log(details)
            if success:
                st.session_state.maint_success_msg = f"儲存成功！ {message}"
                st.cache_data.clear()
                # 【修改點 3】 加入 add_m_contact_date 以便重置
                keys_to_clear = ["add_m_dorm", "add_m_equip", "add_m_date", "add_m_status", "add_m_cat", "add_m_cost", "add_m_vendor", "add_m_contact_date", "add_m_payer", "add_m_finish", "add_m_paid_check", "add_m_desc", "add_m_uploader", "add_m_reporter", "add_m_key_info", "add_m_invoice", "add_m_notes", "add_m_cat_custom"]
                for k in keys_to_clear:
                    if k in st.session_state: del st.session_state[k]
                st.rerun()
            else:
                st.error(message)

def render_progress_tracking(dorm_options, vendor_options, item_type_options, status_options, default_dorm_filter=None):
    """渲染：進度追蹤 (混合模式：快速表格 + 完整詳細編輯)"""
    st.subheader("⏳ 進度追蹤 (未完成案件)")
    
    # --- Part 1: 資料準備與篩選 ---
    raw_df = maintenance_model.get_unfinished_maintenance_logs()
    if raw_df.empty:
        st.success("🎉 恭喜！目前所有維修案件皆已完成。")
        return

    # ID mapping
    dorm_map_id_to_name = dorm_options
    dorm_map_name_to_id = {v: k for k, v in dorm_options.items()}
    raw_df['宿舍'] = raw_df['dorm_id'].map(dorm_map_id_to_name)
    
    vendor_map_id_to_name = vendor_options
    vendor_map_name_to_id = {v: k for k, v in vendor_options.items()}
    raw_df['vendor_id'] = raw_df['vendor_id'].replace({np.nan: None})
    raw_df['廠商'] = raw_df['vendor_id'].map(lambda x: vendor_map_id_to_name.get(x, None))

    with st.expander("🔍 點此展開/收合 篩選條件", expanded=True):
        f_col1, f_col2, f_col3 = st.columns(3)
        available_dorms = sorted(list(raw_df['宿舍'].dropna().unique()))
        available_vendors = sorted(list(raw_df['廠商'].dropna().unique()))
        available_types = sorted(list(raw_df['項目類型'].dropna().unique()))

        # 【核心修改】：設定預設篩選值
        # 如果 URL 傳來 target_dorm，且該宿舍在目前的可用列表中，就預選它
        pre_selected = []
        if default_dorm_filter and default_dorm_filter in available_dorms:
            pre_selected = [default_dorm_filter]

        selected_dorms = f_col1.multiselect("🏠 宿舍地址", options=available_dorms, default=pre_selected, placeholder="全部")
        selected_vendors = f_col2.multiselect("🛠️ 維修廠商", options=available_vendors, placeholder="全部")
        selected_types = f_col3.multiselect("📋 項目類型", options=available_types, placeholder="全部")

    filtered_df = raw_df.copy()
    if selected_dorms: filtered_df = filtered_df[filtered_df['宿舍'].isin(selected_dorms)]
    if selected_vendors: filtered_df = filtered_df[filtered_df['廠商'].isin(selected_vendors)]
    if selected_types: filtered_df = filtered_df[filtered_df['項目類型'].isin(selected_types)]

    st.info(f"顯示 {len(filtered_df)} 筆資料。上方表格可快速編輯；如需**上傳檔案/刪檔/改設備**，請於表格下方選擇案件進行編輯。")

    # --- Part 2: 快速編輯表格 (Data Editor) ---
    display_df = filtered_df[['id', '狀態', '宿舍', '項目類型', '細項說明', '廠商', '通報日期', '提報人']].copy()
    
    edited_df = st.data_editor(
        display_df,
        key="progress_tracking_editor_full",
        use_container_width=True,
        hide_index=True,
        column_config={
            "id": st.column_config.NumberColumn("ID", disabled=True, width="small"),
            "狀態": st.column_config.SelectboxColumn("狀態", options=status_options, required=True),
            "宿舍": st.column_config.SelectboxColumn("宿舍", options=list(dorm_map_name_to_id.keys()), required=True),
            "項目類型": st.column_config.SelectboxColumn("項目類型", options=item_type_options),
            "細項說明": st.column_config.TextColumn("細項說明", width="large"),
            "廠商": st.column_config.SelectboxColumn("廠商", options=list(vendor_map_name_to_id.keys())),
            "通報日期": st.column_config.DateColumn("通報日期", format="YYYY-MM-DD"),
            "提報人": st.column_config.TextColumn("提報人"),
        }
    )

    if st.button("💾 儲存表格變更 (批次)", type="primary", key="btn_save_batch_progress"):
        updates = []
        original_df_idx = display_df.set_index("id")
        edited_df_idx = edited_df.set_index("id")
        for log_id, row in edited_df_idx.iterrows():
            if log_id not in original_df_idx.index: continue
            orig_row = original_df_idx.loc[log_id]
            changes = {}
            if row['狀態'] != orig_row['狀態']: changes['status'] = row['狀態']
            if row['宿舍'] != orig_row['宿舍']:
                if nid := dorm_map_name_to_id.get(row['宿舍']): changes['dorm_id'] = nid
            if row['廠商'] != orig_row['廠商']:
                changes['vendor_id'] = vendor_map_name_to_id.get(row['廠商'])
            if row['項目類型'] != orig_row['項目類型']: changes['item_type'] = row['項目類型']
            if row['細項說明'] != orig_row['細項說明']: changes['description'] = row['細項說明']
            
            new_date = pd.to_datetime(row['通報日期']).date() if pd.notnull(row['通報日期']) else None
            orig_date = pd.to_datetime(orig_row['通報日期']).date() if pd.notnull(orig_row['通報日期']) else None
            if new_date != orig_date: changes['notification_date'] = new_date
            
            if row['提報人'] != orig_row['提報人']: changes['reported_by'] = row['提報人']
            if changes:
                changes['id'] = int(log_id)
                updates.append(changes)
        
        if updates:
            success, msg = maintenance_model.batch_update_logs_all_fields(updates)
            if success: st.success(f"{msg}！"); st.cache_data.clear(); st.rerun()
            else: st.error(msg)
        else:
            st.info("沒有偵測到任何變更。")

    st.markdown("---")
    
    # --- Part 3: 完整詳細編輯區 (檔案/設備/費用) ---
    st.subheader("🛠️ 進階編輯 / 檔案管理")
    
    # 製作選單選項
    # 根據上面的 filtered_df 來製作選單，方便使用者找剛看完的那幾筆
    full_edit_options = {
        row['id']: f"[ID:{row['id']}] {row['宿舍']} - {row['細項說明'][:20]}..." 
        for _, row in filtered_df.sort_values('id', ascending=False).iterrows()
    }
    
    selected_full_edit_id = st.selectbox(
        "請選擇要進行完整編輯 (含上傳檔案) 的案件：",
        options=[None] + list(full_edit_options.keys()),
        format_func=lambda x: "請選擇..." if x is None else full_edit_options.get(x),
        key="sel_full_edit_progress"
    )

    if selected_full_edit_id:
        st.info("👇 下方表單可編輯所有欄位，包含上傳圖片與刪除舊檔。")
        # 呼叫共用元件
        _render_full_edit_form(
            selected_full_edit_id, 
            dorm_options, 
            vendor_options, 
            item_type_options, 
            status_options,
            key_suffix="progress"
        )

def render_edit_delete(dorm_options, vendor_options, item_type_options, status_options):
    """渲染：編輯與刪除 (搜尋模式)"""
    st.subheader("✏️ 編輯 / 刪除單筆維修紀錄 (歷史查詢)")
    all_logs_df = get_all_logs_for_selection()

    if all_logs_df.empty:
        st.info("目前沒有任何維修紀錄。")
        return
    
    search_key = st.text_input("輸入關鍵字搜尋紀錄 (ID, 地址, 說明...)", key="maint_log_search_key")
    filtered_search_df = all_logs_df.copy()

    if search_key:
        keywords = search_key.lower().split()
        filtered_search_df['searchable_text'] = (
            filtered_search_df['id'].astype(str) + " " +
            filtered_search_df['宿舍地址'] + " " +
            filtered_search_df['細項說明'] + " " +
            filtered_search_df['狀態'] + " " +
            filtered_search_df['內部提報人'].fillna('') + " " +
            filtered_search_df['維修廠商'].fillna('') + " " +
            filtered_search_df['項目類型'].fillna('') 
        ).str.lower()
        mask = filtered_search_df['searchable_text'].apply(lambda x: all(k in x for k in keywords))
        filtered_search_df = filtered_search_df[mask].copy()
    
    if filtered_search_df.empty:
         st.warning(f"找不到符合「{search_key}」的紀錄。")
         selected_log_id = None
    else:
        filtered_search_df['通報日期'] = pd.to_datetime(filtered_search_df['通報日期'])
        filtered_search_df = filtered_search_df.sort_values(by=['通報日期', 'id'], ascending=[False, False])
        options_dict = {
            row['id']: (f"[ID:{row['id']}] {row['狀態']} / {row['宿舍地址']} {row['細項說明']} / {row['通報日期'].strftime('%Y-%m-%d')}")
            for _, row in filtered_search_df.iterrows()
        }
        selected_log_id = st.selectbox(
            f"選擇紀錄 (共 {len(filtered_search_df)} 筆)", 
            options=[None] + list(options_dict.keys()), 
            format_func=lambda x: "請選擇..." if x is None else options_dict.get(x), 
            key="selectbox_log_selection"
        )

    if selected_log_id:
        # 呼叫共用元件
        _render_full_edit_form(
            selected_log_id, 
            dorm_options, 
            vendor_options, 
            item_type_options, 
            status_options,
            key_suffix="search"
        )
        
        # 額外功能按鈕 (結案/轉費用/刪除)
        details = maintenance_model.get_single_log_details(selected_log_id)
        if details:
            st.write("")
            c_extra1, c_extra2, c_extra3 = st.columns(3)
            with c_extra1:
                if details.get('status') == '待付款':
                    if st.button("✓ 結案 (已付款)", key=f"btn_complete_{selected_log_id}"):
                        maintenance_model.mark_as_paid_and_complete(selected_log_id)
                        st.cache_data.clear()
                        st.rerun()
            with c_extra2:
                if not details.get('is_archived_as_expense') and details.get('status') in ['待付款', '已完成'] and (details.get('cost') or 0) > 0 and details.get('payer') == '我司':
                    if st.button("💰 轉入年度費用", key=f"btn_archive_{selected_log_id}"):
                        maintenance_model.archive_log_as_annual_expense(selected_log_id)
                        st.cache_data.clear()
                        st.rerun()
            with c_extra3:
                 if st.button("🗑️ 刪除紀錄", key=f"btn_del_{selected_log_id}", type="primary"):
                     maintenance_model.delete_log(selected_log_id)
                     st.cache_data.clear()
                     st.rerun()

def render_overview(dorm_options, vendor_options, item_type_options, status_options):
    """渲染：維修紀錄總覽 (新增下方編輯功能)"""
    st.subheader("📊 維修紀錄總覽")
    
    # --- Part 1: 篩選條件 ---
    c1, c2, c3 = st.columns(3)
    f_status = c1.selectbox("狀態篩選", [""] + status_options, key="ov_status")
    f_dorm = c2.selectbox("宿舍篩選", [None] + list(dorm_options.keys()), format_func=lambda x: "全部" if x is None else dorm_options.get(x), key="ov_dorm")
    f_vendor = c3.selectbox("廠商篩選", [None] + list(vendor_options.keys()), format_func=lambda x: "全部" if x is None else vendor_options.get(x), key="ov_vendor")
    
    c4, c5 = st.columns(2)
    f_start = c4.date_input("完成日期 (起)", value=None, key="ov_start")
    f_end = c5.date_input("完成日期 (迄)", value=None, key="ov_end")

    filters = {}
    if f_status: filters["status"] = f_status
    if f_dorm: filters["dorm_id"] = f_dorm
    if f_vendor: filters["vendor_id"] = f_vendor
    if f_start: filters["start_date"] = f_start
    if f_end: filters["end_date"] = f_end

    # 取得篩選後的資料
    log_df = maintenance_model.get_logs_for_view(filters)
    
    # --- Part 2: 顯示資料表格 ---
    if not log_df.empty:
        if f_vendor or f_start or f_end:
             st.success(f"篩選總計: {len(log_df)} 筆, 費用總額: NT$ {log_df['維修費用'].sum():,}")
        
        # 顯示表格 (增加搜尋功能以便快速定位 ID)
        st.dataframe(log_df, width='stretch', hide_index=True)
        
        st.markdown("---")
        
        # --- Part 3: 下方編輯區 (新增) ---
        st.subheader("🔍 檢視詳情與編輯")
        st.info("請在下方選擇上方表格中的案件 ID，即可進行詳細內容編輯或查看照片。")
        
        # 準備下拉選單選項
        edit_options = {
            row['id']: f"[ID:{row['id']}] {row['宿舍地址']} - {row['細項說明'][:30]}..." 
            for _, row in log_df.iterrows()
        }
        
        selected_log_id = st.selectbox(
            "選擇要編輯/檢視的案件：",
            options=[None] + list(edit_options.keys()),
            format_func=lambda x: "請選擇案件 ID..." if x is None else edit_options.get(x),
            key="overview_edit_selector"
        )
        
        if selected_log_id:
            # 呼叫現有的共用編輯表單組件
            _render_full_edit_form(
                selected_log_id, 
                dorm_options, 
                vendor_options, 
                item_type_options, 
                status_options,
                key_suffix="overview"
            )
    else:
        st.info("無符合條件的資料")

def render_batch_archive():
    """渲染：批次轉入年度費用"""
    st.subheader("📦 批次轉入年度費用")
    st.info("列出已完成/待付款且為「我司」支付，但尚未歸檔的項目。")

    @st.cache_data
    def get_archivable_data():
        return maintenance_model.get_archivable_logs()

    archivable_df = get_archivable_data()

    if archivable_df.empty:
        st.success("目前沒有可批次轉入的項目。")
        return

    if 'maint_archive_default' not in st.session_state: st.session_state.maint_archive_default = False
    if 'maint_archive_reset' not in st.session_state: st.session_state.maint_archive_reset = 0

    c_tools1, c_tools2 = st.columns(2)
    if c_tools1.button("✅ 全選"):
        st.session_state.maint_archive_default = True
        st.session_state.maint_archive_reset += 1
        st.rerun()
    if c_tools2.button("⬜ 取消全選"):
        st.session_state.maint_archive_default = False
        st.session_state.maint_archive_reset += 1
        st.rerun()

    df_with_select = archivable_df.copy()
    df_with_select.insert(0, "選取", st.session_state.maint_archive_default)
    
    edited_df = st.data_editor(
        df_with_select,
        hide_index=True,
        column_config={"選取": st.column_config.CheckboxColumn(required=True)},
        disabled=archivable_df.columns,
        key=f"archive_editor_{st.session_state.maint_archive_reset}"
    )
    
    selected_rows = edited_df[edited_df.選取]
    
    if st.button("🚀 執行批次轉入", type="primary", disabled=selected_rows.empty):
        ids = selected_rows['id'].tolist()
        with st.spinner(f"處理 {len(ids)} 筆資料..."):
            s_count, f_count = maintenance_model.batch_archive_logs(ids)
        
        if s_count: st.success(f"成功轉入 {s_count} 筆！")
        if f_count: st.error(f"失敗 {f_count} 筆。")
        
        st.session_state.maint_archive_default = False
        st.session_state.maint_archive_reset += 1
        st.cache_data.clear()
        st.rerun()

# -----------------------------------------------------------------------------
# 新增改善建議 (連續輸入模式)
# -----------------------------------------------------------------------------
def render_add_improvement_suggestion(dorm_options, vendor_options, item_type_options, status_options):
    """渲染：新增改善建議 (支援選定宿舍後連續新增，並可選房號)"""
    st.subheader("💡 新增改善建議 / 待辦事項")
    st.info("此模式用於快速建立「待改善」項目。選擇宿舍後，可連續輸入多筆資料，無需重複選擇宿舍。")

    # --- 1. 宿舍選擇 (放在 Form 外層以保持狀態) ---
    c1, c2 = st.columns([1, 2])
    with c1:
        dorm_keys = list(dorm_options.keys())
        selected_dorm_id = st.selectbox(
            "請先選擇宿舍*", 
            options=dorm_keys, 
            format_func=lambda x: dorm_options.get(x, "未選擇"), 
            key="improve_dorm_select"
        )
    
    # 根據選到的宿舍抓取房號資料
    room_options = {}
    if selected_dorm_id:
        rooms_df = maintenance_model.get_rooms_for_selector(selected_dorm_id)
        if not rooms_df.empty:
            room_options = {row['id']: row['room_number'] for _, row in rooms_df.iterrows()}

    # --- 2. 輸入表單 ---
    with st.form("improvement_form", clear_on_submit=True):
        st.markdown(f"**正在新增：{dorm_options.get(selected_dorm_id, '')}**")
        
        # 【修改】改為 4 欄，加入日期選擇
        c_row1_1, c_row1_2, c_row1_3, c_row1_4 = st.columns(4)
        
        with c_row1_1:
            # 新增：通報日期 (預設為今天)
            new_notify_date = st.date_input("通報日期*", value=date.today())

        with c_row1_2:
            # 房號選擇
            r_keys = [None] + list(room_options.keys())
            selected_room_id = st.selectbox(
                "房號 (選填)", 
                options=r_keys, 
                format_func=lambda x: "公共區域" if x is None else f"{room_options.get(x)}",
            )
        
        with c_row1_3:
            # 狀態
            default_status = "待改善"
            st_idx = status_options.index(default_status) if default_status in status_options else 0
            new_status = st.selectbox("案件狀態*", options=status_options, index=st_idx)
            
        with c_row1_4:
            # 類型
            new_category_sel = st.selectbox("項目類別", options=item_type_options)

        # 說明與備註
        new_description = st.text_area("改善建議 / 缺失說明*", height=100, placeholder="請描述需要改善的項目...")
        
        c_row2_1, c_row2_2 = st.columns(2)
        new_reporter = c_row2_1.text_input("提報人", placeholder="您的姓名")
        new_notes = c_row2_2.text_input("備註 (可選)", placeholder="例如: 急件")

        # 照片上傳
        uploaded_files = st.file_uploader("📷 上傳照片 (可多選)", type=['jpg', 'jpeg', 'png', 'pdf'], accept_multiple_files=True)

        # 送出按鈕
        submit_btn = st.form_submit_button("💾 儲存並新增下一筆", type="primary")

        if submit_btn:
            if not selected_dorm_id or not new_description:
                st.error("「宿舍」和「說明」為必填欄位！")
            else:
                final_description = new_description
                if selected_room_id:
                    room_num = room_options.get(selected_room_id)
                    final_description = f"【房號: {room_num}】 {new_description}"

                file_paths = []
                if uploaded_files:
                    file_info_dict = {
                        "date": new_notify_date.strftime('%Y%m%d'), # 使用選取的日期
                        "address": dorm_options.get(selected_dorm_id, 'Unknown'), 
                        "reporter": new_reporter, 
                        "type": new_category_sel
                    }
                    for file in uploaded_files:
                        path = maintenance_model.save_uploaded_photo(file, file_info_dict)
                        file_paths.append(path)
                
                details = {
                    'dorm_id': selected_dorm_id, 
                    'equipment_id': None, 
                    'vendor_id': None, 
                    'status': new_status,
                    'notification_date': new_notify_date, # 【修改】使用選取的日期
                    'reported_by': new_reporter, 
                    'item_type': new_category_sel, 
                    'description': final_description,
                    'contacted_vendor_date': None, 
                    'completion_date': None, 
                    'key_info': "",    
                    'cost': 0, 
                    'payer': "", 
                    'invoice_date': None, 
                    'invoice_info': "", 
                    'notes': new_notes, 
                    'photo_paths': file_paths 
                }
                
                success, message = maintenance_model.add_log(details)
                
                if success:
                    st.success(f"✅ 已新增：{final_description}")
                    st.cache_data.clear()
                else:
                    st.error(message)

    if selected_dorm_id:
        st.markdown("---")
        st.caption(f"📋 {dorm_options.get(selected_dorm_id)} - 近期新增項目：")
        logs = maintenance_model.get_logs_for_view({"dorm_id": selected_dorm_id, "status": "待改善"})
        if not logs.empty:
            # 這裡也可以多顯示通報日期
            st.dataframe(
                logs[['id', '通報日期', '細項說明', '狀態']].head(5), 
                hide_index=True,
                use_container_width=True
            )

# -----------------------------------------------------------------------------
# 渲染匯出報表頁面
# -----------------------------------------------------------------------------
def render_export_report(dorm_options, status_options):
    st.subheader("📑 匯出改善/維修報表")
    st.info("支援匯出 Excel (連結用)、HTML (檢視用) 或 Word (現場簽核/手寫用)。")

    # --- 1. 篩選條件 ---
    c1, c2 = st.columns(2)
    selected_dorm_ids = c1.multiselect(
        "1. 選擇宿舍 (留空則全選)", 
        options=list(dorm_options.keys()), 
        format_func=lambda x: dorm_options.get(x)
    )
    
    default_statuses = [s for s in ["待改善", "待處理"] if s in status_options]
    selected_statuses = c2.multiselect(
        "2. 選擇案件狀態/類型 (留空則全選)",
        options=status_options,
        default=default_statuses,
        key="export_status_filter"
    )

    df = maintenance_model.get_logs_for_view(filters=None)
    
    if not df.empty:
        if selected_dorm_ids:
            df = df[df['dorm_id'].isin(selected_dorm_ids)]
        if selected_statuses:
            df = df[df['狀態'].isin(selected_statuses)]

    if df.empty:
        st.warning("⚠️ 目前條件下無任何資料。")
        return

    # --- 排序邏輯 ---
    df['sort_room'] = df['細項說明'].apply(extract_room_number)
    df = df.sort_values(by=['宿舍地址', '項目類型', 'sort_room', '通報日期'], ascending=[True, True, True, False])

    st.success(f"已篩選出 {len(df)} 筆資料。")

    # === 【新增】檔案名稱產生邏輯 ===
    today_str = date.today().strftime('%Y%m%d')
    # 預設檔名 (多選或未選特定宿舍時)
    file_name_base = f"維修報表_彙總_{today_str}"
    
    # 若只選擇了一間宿舍，則依照要求格式化檔名: 編號_地址_日期
    if selected_dorm_ids and len(selected_dorm_ids) == 1:
        did = selected_dorm_ids[0]
        # 從資料庫取得詳細資料 (確保有編號)
        d_info = dormitory_model.get_dorm_details_by_id(did)
        if d_info:
            d_code = d_info.get('legacy_dorm_code') or "無編號"
            d_addr = d_info.get('original_address') or "未知地址"
            
            # 移除檔名不允許的字元 (如 / \ 等)
            safe_code = str(d_code).replace("/", "_").replace("\\", "_").strip()
            safe_addr = str(d_addr).replace("/", "_").replace("\\", "_").strip()
            
            file_name_base = f"{safe_code}_{safe_addr}_{today_str}"
    
    st.caption(f"ℹ️ 預計存檔名稱：`{file_name_base}.xxx` (實際存檔位置請於跳出的視窗中選擇)")
    # =================================

    col_dl1, col_dl2, col_dl3 = st.columns(3)

    # === 選項 A: Excel 報表 ===
    with col_dl1:
        if st.button("📥 Excel 報表 (含連結)", use_container_width=True):
            export_df = pd.DataFrame()
            export_df['地址'] = df['宿舍地址']
            export_df['房號'] = df['細項說明'].apply(extract_room_number)
            export_df['通報日期'] = df['通報日期']
            export_df['項目類別'] = df['項目類型']
            export_df['改善建議/缺失說明'] = df['細項說明']
            export_df['案件狀態'] = df['狀態']
            export_df['照片'] = df['photo_paths'].apply(lambda x: f"有 ({len(x)})" if x and len(x) > 0 else "")

            # 建立連結
            base_url = "http://192.168.1.116:8501/" 
            page_param = "page=" + quote("維修追蹤與費用")
            
            def create_link(row):
                dorm_name = row['地址']
                full_url = f"{base_url}?{page_param}&view_mode=tracking&target_dorm={quote(dorm_name)}"
                return full_url

            export_df['系統連結'] = export_df.apply(create_link, axis=1)

            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                export_df.to_excel(writer, index=False, sheet_name='報表')
                workbook = writer.book
                worksheet = writer.sheets['報表']
                link_format = workbook.add_format({'font_color': 'blue', 'underline': 1})
                worksheet.set_column('A:A', 20)
                worksheet.set_column('B:B', 8)
                worksheet.set_column('C:C', 12)
                worksheet.set_column('E:E', 45)
                
                col_idx = 7
                for row_idx, url in enumerate(export_df['系統連結']):
                    if url:
                        worksheet.write_url(row_idx + 1, col_idx, url, link_format, string='前往系統處理')

            st.download_button(
                label="📄 下載 Excel",
                data=output.getvalue(),
                file_name=f"{file_name_base}.xlsx", # 套用新檔名
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="btn_dl_excel"
            )

    # === 選項 B: HTML 報表 ===
    with col_dl2:
        if st.button("🖼️ HTML 報表 (含照片)", use_container_width=True):
             # 建立 HTML 字串
            html_content = f"""
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: "Microsoft JhengHei", Arial, sans-serif; padding: 20px; }}
                    h2 {{ color: #333; }}
                    table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; vertical-align: top; }}
                    th {{ background-color: #f2f2f2; }}
                    .photo-cell img {{ max-width: 150px; max-height: 150px; margin: 2px; border: 1px solid #ccc; }}
                    .status-tag {{ font-weight: bold; color: #d9534f; }}
                </style>
            </head>
            <body>
                <h2>🔍 改善/維修建議報表</h2>
                <p>匯出日期: {date.today()} | 篩選條件: {", ".join(selected_statuses) if selected_statuses else "全部"}</p>
                <table>
                    <thead>
                        <tr>
                            <th style="width: 12%">地址 / 房號</th>
                            <th style="width: 10%">通報日期</th>
                            <th style="width: 10%">類別</th>
                            <th style="width: 10%">狀態</th>
                            <th style="width: 30%">說明</th>
                            <th style="width: 28%">照片證明</th>
                        </tr>
                    </thead>
                    <tbody>
            """
            
            for _, row in df.iterrows():
                dorm_addr = row['宿舍地址']
                room_no = extract_room_number(row['細項說明'])
                addr_display = f"{dorm_addr}<br><b>{room_no}</b>" if room_no else dorm_addr
                
                photos_html = ""
                if row['photo_paths'] and isinstance(row['photo_paths'], list):
                    for path in row['photo_paths']:
                        b64_str = image_to_base64(path)
                        if b64_str:
                            photos_html += f'<a href="{b64_str}" target="_blank"><img src="{b64_str}"></a>'
                
                if not photos_html: photos_html = "<span style='color:#ccc'>無照片</span>"

                html_content += f"""
                        <tr>
                            <td>{addr_display}</td>
                            <td>{row['通報日期']}</td>
                            <td>{row['項目類型']}</td>
                            <td><span class="status-tag">{row['狀態']}</span></td>
                            <td>{row['細項說明']}</td>
                            <td class="photo-cell">{photos_html}</td>
                        </tr>
                """
            
            html_content += """
                    </tbody>
                </table>
            </body>
            </html>
            """
            
            st.download_button(
                label="📄 下載 HTML",
                data=html_content,
                file_name=f"{file_name_base}.html", # 套用新檔名
                mime="text/html",
                key="btn_dl_html"
            )

    # === 選項 C: Word 簽核單 (優化排版版) ===
    with col_dl3:
        if st.button("📝 Word 簽核單 (手寫用)", type="primary", use_container_width=True):
            try:
                # 建立文件
                doc = Document()
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11.69) 
                section.page_height = Inches(8.27) 
                
                section.left_margin = Cm(1.0)
                section.right_margin = Cm(1.0)
                section.top_margin = Cm(1.0)
                section.bottom_margin = Cm(1.0)

                # 標題
                heading = doc.add_heading(f'宿舍改善/維修執行簽核單', 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 建立表格
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                
                hdr_tr = table.rows[0]._tr
                hdr_trPr = hdr_tr.get_or_add_trPr()
                tblHeader = OxmlElement('w:tblHeader')
                hdr_trPr.append(tblHeader)

                # 設定表頭內容
                hdr_cells = table.rows[0].cells
                headers = ["地址/房號", "日期", "類別", "說明", "現況照片", "改善時間", "改善說明", "改善後相片"]
                
                for i, h in enumerate(headers):
                    set_cell_font(hdr_cells[i], h, bold=True, font_size=11)

                # 填入資料
                for _, row in df.iterrows():
                    new_row = table.add_row()
                    
                    tr = new_row._tr
                    trPr = tr.get_or_add_trPr()
                    cantSplit = OxmlElement('w:cantSplit')
                    trPr.append(cantSplit)
                    
                    row_cells = new_row.cells
                    
                    # 1. 地址/房號
                    room_no = extract_room_number(row['細項說明'])
                    addr_text = f"{row['宿舍地址']}\n{room_no}" if room_no else row['宿舍地址']
                    set_cell_font(row_cells[0], addr_text)
                    
                    # 2. 日期
                    set_cell_font(row_cells[1], str(row['通報日期']))
                    
                    # 3. 類別
                    set_cell_font(row_cells[2], row['項目類型'])
                    
                    # 4. 說明
                    desc_clean = re.sub(r"【房號: .+?】", "", row['細項說明']).strip()
                    set_cell_font(row_cells[3], desc_clean)
                    
                    # 5. 現況照片
                    photo_paragraph = row_cells[4].paragraphs[0]
                    if row['photo_paths'] and isinstance(row['photo_paths'], list):
                        count = 0
                        for path in row['photo_paths']:
                            if os.path.exists(path) and count < 2:
                                try:
                                    run = photo_paragraph.add_run()
                                    run.add_picture(path, width=Cm(3.5))
                                    run.add_text("\n")
                                    count += 1
                                except Exception:
                                    pass
                        if count == 0:
                            photo_paragraph.add_run("無照片")
                    else:
                        photo_paragraph.add_run("無照片")

                    # 6, 7, 8 留空欄位
                    set_cell_font(row_cells[5], "\n\n\n") 
                    set_cell_font(row_cells[6], "") 
                    set_cell_font(row_cells[7], "") 

                # 儲存
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                st.download_button(
                    label="📄 下載 Word",
                    data=doc_io,
                    file_name=f"{file_name_base}.docx", # 套用新檔名
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_dl_word"
                )

            except NameError:
                st.error("請先安裝 python-docx 套件： pip install python-docx")
            except Exception as e:
                st.error(f"Word 匯出失敗: {e}")

def render():
    # 處理 URL 參數
    query_params = st.query_params
    default_nav_mode = "➕ 新增維修紀錄"
    tracking_target_dorm = None

    if query_params.get("view_mode") == "tracking":
        default_nav_mode = "⏳ 未完成案件追蹤"
        tracking_target_dorm = query_params.get("target_dorm")

    st.header("維修追蹤管理")
    st.info("用於登記、追蹤和管理宿舍的各項維修申報與進度，並可上傳現場照片、報價單(PDF)等相關文件。")
    
    dorms = dormitory_model.get_dorms_for_selection()
    dorm_options = {d['id']: f"({d.get('legacy_dorm_code') or '無編號'}) {d.get('original_address', '')}" for d in dorms} if dorms else {}
    
    vendors = vendor_model.get_vendors_for_view()
    vendor_options = {v['id']: f"{v['服務項目']} - {v['廠商名稱']}" for _, v in vendors.iterrows()} if not vendors.empty else {}
    
    status_options = ["待處理", "待改善", "待尋廠商", "進行中", "待付款", "已完成"]
    item_type_options = ["維修", "定期保養", "更換耗材", "水電", "包通", "飲水機", "冷氣", "消防", "金城", "監視器", "水質檢測", "清運", "裝潢", "油漆", "蝦皮", "泥作", "宣導", "其他(手動輸入)"]

    nav_options = ["➕ 新增維修紀錄", "➕ 新增改善建議", "⏳ 未完成案件追蹤", "✏️ 編輯 / 刪除單筆維修紀錄", "📊 維修紀錄總覽", "📦 批次轉入年度費用", "📑 匯出改善/維修報表"]
    try:
        default_idx = nav_options.index(default_nav_mode)
    except ValueError:
        default_idx = 0

    app_mode = st.radio(
        "請選擇操作項目：",
        nav_options,
        index=default_idx,
        horizontal=True,
        key="maintenance_main_nav"
    )
    st.markdown("---")

    if app_mode == "➕ 新增維修紀錄":
        render_add_new_record(dorm_options, vendor_options, item_type_options, status_options)
    elif app_mode == "➕ 新增改善建議":
        if 'render_add_improvement_suggestion' in globals():
            render_add_improvement_suggestion(dorm_options, vendor_options, item_type_options, status_options)
        else:
            st.warning("請確認程式碼是否包含 render_add_improvement_suggestion 函式")
    elif app_mode == "⏳ 未完成案件追蹤":
        render_progress_tracking(dorm_options, vendor_options, item_type_options, status_options, default_dorm_filter=tracking_target_dorm)
    elif app_mode == "✏️ 編輯 / 刪除單筆維修紀錄":
        render_edit_delete(dorm_options, vendor_options, item_type_options, status_options)
    elif app_mode == "📊 維修紀錄總覽":
        render_overview(dorm_options, vendor_options, item_type_options, status_options) # 補上後兩個參數
    elif app_mode == "📦 批次轉入年度費用":
        render_batch_archive()
    elif app_mode == "📑 匯出改善/維修報表":
        # 【修正】這裡必須傳入 status_options
        render_export_report(dorm_options, status_options)